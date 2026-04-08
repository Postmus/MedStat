import json
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUTPUT_DIR = Path(__file__).resolve().parent
DATA_FILENAME = "certificate_data.json"
BACKGROUND_FILENAME = "background_logo.jpg"
DOCX_OUTPUT_FILENAME = "Certificate_MedStat_editable.docx"
PDF_OUTPUT_FILENAME = "Certificate_MedStat.pdf"
PNG_OUTPUT_FILENAME = "Certificate_MedStat_preview.png"

PAGE_WIDTH_PT = 595.32001
PAGE_HEIGHT_PT = 841.92004
PAGE_WIDTH_PX = 2480
PAGE_HEIGHT_PX = 3508
PX_PER_PT = PAGE_WIDTH_PX / PAGE_WIDTH_PT

BACKGROUND_X_PT = 70.839084
BACKGROUND_Y_PT = 107.837662
BACKGROUND_SCALE_X = 2.044094
BACKGROUND_SCALE_Y = 2.050391
BACKGROUND_BASE_WIDTH_PX = 224
BACKGROUND_BASE_HEIGHT_PX = 311

OVERLAY_X_PT = 25.890625
OVERLAY_Y_PT = 101.363281
OVERLAY_W_PT = 565.273438 - 25.890625
OVERLAY_H_PT = 793.574219 - 101.363281

TITLE_CENTER_X_PT = (88.199999 + 539.157083) / 2
TITLE_TOP_Y_PT = 121.366426
COURSE_CENTER_X_PT = (146.759995 + 480.468248) / 2
COURSE_TOP_Y_PT = 238.511002
PERIOD_CENTER_X_PT = (218.759995 + 408.264376) / 2
PERIOD_TOP_Y_PT = 271.851425
ECTS_CENTER_X_PT = (248.160004 + 379.020325) / 2
ECTS_TOP_Y_PT = 294.190994
NAME_CENTER_X_PT = (216.120003 + 375.553773) / 2
NAME_TOP_Y_PT = 382.390999
ISSUE_X_PT = 103.919999
ISSUE_TOP_Y_PT = 552.815478
SIGNATURE_LINE_X_PT = 105.122705
SIGNATURE_LINE_TOP_Y_PT = 619.775499
SIGNATURE_NAME_X_PT = 105.122705
SIGNATURE_NAME_TOP_Y_PT = 636.455492
FOOTER_X_PT = 103.919999
FOOTER_TOP_Y_PT = 720.095507

BOLD_FONT = "/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf"
REGULAR_FONT = "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf"


def set_page_background(section, color_hex: str) -> None:
    sect_pr = section._sectPr
    pg_borders = sect_pr.first_child_found_in("w:pgBorders")
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        pg_borders.set(qn("w:offsetFrom"), "page")
        sect_pr.append(pg_borders)

    for edge in ("top", "left", "bottom", "right"):
        border = pg_borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            pg_borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "18")
        border.set(qn("w:space"), "24")
        border.set(qn("w:color"), color_hex)


def add_paragraph(document, text, *, size, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=0):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    return paragraph


def safe_stem(value: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9]+", "_", value.strip()).strip("_")
    return stem or "student"


def build_document() -> Document:
    data = load_data()
    student_name = data["student_name"]
    course_period = data["course_period"]
    issue_city_and_date = data["issue_city_and_date"]

    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.6)
    section.bottom_margin = Cm(2.6)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)
    set_page_background(section, "D9D9D9")

    add_paragraph(document, "", size=1, spacing_after=28)
    add_paragraph(document, "Certificate of attendance", size=23, bold=True, spacing_after=14)
    add_paragraph(document, "Course on Medical Statistics", size=16, bold=True, spacing_after=2)
    add_paragraph(document, course_period, size=12, bold=True, spacing_after=6)
    add_paragraph(document, "(2,5 ECTS)", size=16, bold=True, spacing_after=58)

    name_paragraph = add_paragraph(document, student_name, size=18, bold=True, spacing_after=74)
    for run in name_paragraph.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    add_paragraph(
        document,
        issue_city_and_date,
        size=11,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        spacing_after=36,
    )

    line = add_paragraph(document, "_________________________", size=12, align=WD_ALIGN_PARAGRAPH.LEFT, spacing_after=10)
    for run in line.runs:
        run.font.name = "Courier New"

    add_paragraph(document, "Dr. D. Postmus", size=11, align=WD_ALIGN_PARAGRAPH.LEFT, spacing_after=40)

    footer = add_paragraph(
        document,
        "Department of Epidemiology\nUniversity Medical Center Groningen\nUniversity of Groningen",
        size=11,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        spacing_after=0,
    )
    footer.paragraph_format.line_spacing = 1.0

    return document


def load_data() -> dict:
    data_path = OUTPUT_DIR / DATA_FILENAME
    with data_path.open("r", encoding="utf-8") as handle:
        data = json.load(handle)

    required_fields = ("student_name", "course_period", "issue_city_and_date")
    missing = [field for field in required_fields if not data.get(field)]
    if missing:
        missing_text = ", ".join(missing)
        raise ValueError(f"Missing required fields in {data_path.name}: {missing_text}")

    return data


def pt_to_px(value: float) -> int:
    return int(round(value * PX_PER_PT))


def load_font(path: str, size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(path, size=size)


def draw_centered_text(draw, text: str, *, center_x_pt: float, top_y_pt: float, font, fill="black") -> None:
    bbox = draw.textbbox((0, 0), text, font=font)
    width = bbox[2] - bbox[0]
    x = pt_to_px(center_x_pt) - width // 2
    y = pt_to_px(top_y_pt)
    draw.text((x, y), text, font=font, fill=fill)


def draw_left_text(draw, text: str, *, x_pt: float, top_y_pt: float, font, fill="black", line_gap_px: int = 10) -> None:
    x = pt_to_px(x_pt)
    y = pt_to_px(top_y_pt)
    for line in text.splitlines():
        draw.text((x, y), line, font=font, fill=fill)
        bbox = draw.textbbox((x, y), line, font=font)
        y += (bbox[3] - bbox[1]) + line_gap_px


def build_pdf_assets() -> tuple[Image.Image, Path, Path]:
    data = load_data()
    background_path = OUTPUT_DIR / BACKGROUND_FILENAME

    page = Image.new("RGBA", (PAGE_WIDTH_PX, PAGE_HEIGHT_PX), "white")

    background = Image.open(background_path).convert("RGBA")
    background_size = (
        pt_to_px(BACKGROUND_BASE_WIDTH_PX * BACKGROUND_SCALE_X),
        pt_to_px(BACKGROUND_BASE_HEIGHT_PX * BACKGROUND_SCALE_Y),
    )
    background = background.resize(background_size)
    page.paste(background, (pt_to_px(BACKGROUND_X_PT), pt_to_px(BACKGROUND_Y_PT)), background)

    overlay = Image.new(
        "RGBA",
        (pt_to_px(OVERLAY_W_PT), pt_to_px(OVERLAY_H_PT)),
        (255, 255, 255, 128),
    )
    page.alpha_composite(overlay, (pt_to_px(OVERLAY_X_PT), pt_to_px(OVERLAY_Y_PT)))

    draw = ImageDraw.Draw(page)

    title_font = load_font(BOLD_FONT, 170)
    subtitle_font = load_font(BOLD_FONT, 112)
    small_bold_font = load_font(BOLD_FONT, 74)
    name_font = load_font(BOLD_FONT, 112)
    regular_font = load_font(REGULAR_FONT, 56)

    draw_centered_text(draw, "Certificate of attendance", center_x_pt=TITLE_CENTER_X_PT, top_y_pt=TITLE_TOP_Y_PT, font=title_font)
    draw_centered_text(draw, "Course on Medical Statistics", center_x_pt=COURSE_CENTER_X_PT, top_y_pt=COURSE_TOP_Y_PT, font=subtitle_font)
    draw_centered_text(draw, data["course_period"], center_x_pt=PERIOD_CENTER_X_PT, top_y_pt=PERIOD_TOP_Y_PT, font=small_bold_font)
    draw_centered_text(draw, "(2,5 ECTS)", center_x_pt=ECTS_CENTER_X_PT, top_y_pt=ECTS_TOP_Y_PT, font=subtitle_font)
    draw_centered_text(draw, data["student_name"], center_x_pt=NAME_CENTER_X_PT, top_y_pt=NAME_TOP_Y_PT, font=name_font)

    draw_left_text(draw, data["issue_city_and_date"], x_pt=ISSUE_X_PT, top_y_pt=ISSUE_TOP_Y_PT, font=regular_font)
    draw_left_text(draw, "_________________________", x_pt=SIGNATURE_LINE_X_PT, top_y_pt=SIGNATURE_LINE_TOP_Y_PT, font=regular_font)
    draw_left_text(draw, "Dr. D. Postmus", x_pt=SIGNATURE_NAME_X_PT, top_y_pt=SIGNATURE_NAME_TOP_Y_PT, font=regular_font)
    draw_left_text(
        draw,
        "Department of Epidemiology\nUniversity Medical Center Groningen\nUniversity of Groningen",
        x_pt=FOOTER_X_PT,
        top_y_pt=FOOTER_TOP_Y_PT,
        font=regular_font,
        line_gap_px=10,
    )

    pdf_path = OUTPUT_DIR / PDF_OUTPUT_FILENAME
    png_path = OUTPUT_DIR / PNG_OUTPUT_FILENAME
    return page, pdf_path, png_path


def main() -> None:
    data = load_data()
    document = build_document()

    output_path = OUTPUT_DIR / DOCX_OUTPUT_FILENAME
    student_copy_path = OUTPUT_DIR / f"Certificate_MedStat_{safe_stem(data['student_name'])}.docx"
    student_pdf_copy_path = OUTPUT_DIR / f"Certificate_MedStat_{safe_stem(data['student_name'])}.pdf"

    document.save(output_path)
    document.save(student_copy_path)

    page, pdf_path, png_path = build_pdf_assets()
    page_rgb = page.convert("RGB")
    page_rgb.save(pdf_path, resolution=300.0)
    page_rgb.save(student_pdf_copy_path, resolution=300.0)
    page_rgb.save(png_path)

    print(f"Wrote {output_path}")
    print(f"Wrote {student_copy_path}")
    print(f"Wrote {pdf_path}")
    print(f"Wrote {student_pdf_copy_path}")
    print(f"Wrote {png_path}")


if __name__ == "__main__":
    main()
